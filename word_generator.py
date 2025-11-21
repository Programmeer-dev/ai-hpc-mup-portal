from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

def generate_docx_confirmation(user_name, user_email, service_name, service_details, 
                                center_name, center_address):
    """
    Generiše Word dokument sa informacijama o MUP usluzi.
    
    Args:
        user_name (str): Ime korisnika
        user_email (str): Email korisnika
        service_name (str): Naziv usluge
        service_details (dict): Detalji usluge (cijena, trajanje, potrebni dokumenti)
        center_name (str): Naziv MUP centra
        center_address (str): Adresa MUP centra
    
    Returns:
        BytesIO: Word dokument kao bytes objekat
    """
    # Kreiraj novi dokument
    doc = Document()
    
    # Podesi margine
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Naslov
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('INFORMACIJE O MUP USLUZI')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Podnaslov
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Ministarstvo unutrašnjih poslova Crne Gore')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Prazan red
    
    # Datum izdavanja
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'Datum izdavanja: {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.italic = True
    
    doc.add_paragraph()  # Prazan red
    
    # Informacije o korisniku
    heading1 = doc.add_heading('PODACI O KORISNIKU', level=2)
    heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Light Grid Accent 1'
    
    # Red 1
    table1.rows[0].cells[0].text = 'Ime i prezime:'
    table1.rows[0].cells[1].text = user_name
    
    # Red 2
    table1.rows[1].cells[0].text = 'Email:'
    table1.rows[1].cells[1].text = user_email
    
    # Stilizuj prvu kolonu (bold)
    for row in table1.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Prazan red
    
    # Informacije o usluzi
    heading2 = doc.add_heading('PODACI O USLUZI', level=2)
    heading2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Light Grid Accent 1'
    
    # Red 1
    table2.rows[0].cells[0].text = 'Usluga:'
    table2.rows[0].cells[1].text = service_name
    
    # Red 2
    table2.rows[1].cells[0].text = 'Cijena:'
    table2.rows[1].cells[1].text = service_details['cijena']
    
    # Red 3
    table2.rows[2].cells[0].text = 'Vrijeme obrade:'
    table2.rows[2].cells[1].text = service_details['trajanje']
    
    # Stilizuj prvu kolonu (bold)
    for row in table2.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Prazan red
    
    # Potrebni dokumenti
    heading3 = doc.add_heading('POTREBNA DOKUMENTACIJA', level=2)
    heading3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    documents = service_details['potrebni_dokumenti']
    if documents:
        for doc_item in documents:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(doc_item)
    else:
        doc.add_paragraph('Nije navedeno.')
    
    doc.add_paragraph()  # Prazan red
    
    # Lokacija MUP centra
    heading4 = doc.add_heading('LOKACIJA MUP CENTRA', level=2)
    heading4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Light Grid Accent 1'
    
    # Red 1
    table3.rows[0].cells[0].text = 'Centar:'
    table3.rows[0].cells[1].text = center_name
    
    # Red 2
    table3.rows[1].cells[0].text = 'Adresa:'
    table3.rows[1].cells[1].text = center_address
    
    # Stilizuj prvu kolonu (bold)
    for row in table3.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Prazan red
    doc.add_paragraph()  # Prazan red
    
    # Napomena na dnu
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Molimo donesite sve potrebne dokumente kada dođete u MUP centar.')
    footer_run.font.size = Pt(10)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.add_run('Za više informacija kontaktirajte: info@mup.gov.me | +382 20 123 456')
    footer2_run.font.size = Pt(9)
    footer2_run.font.color.rgb = RGBColor(150, 150, 150)
    
    # Sačuvaj dokument u BytesIO objekat
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer
